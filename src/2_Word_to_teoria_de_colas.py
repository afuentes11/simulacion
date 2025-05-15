import re
from docx import Document

# Leer texto desde un archivo Word
def leer_texto_desde_word(ruta_word):
    doc = Document(ruta_word)
    texto = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return texto

# Determinar el tamaño de la población (finita o infinita)
def determinar_tamano_poblacion(texto):
    try:
        # Buscar indicios de población finita
        mencion_poblacion_finita = re.search(r'(la población|los clientes|los usuarios) (es|son) finita(s)?', texto, re.IGNORECASE)
        if mencion_poblacion_finita:
            return "finita", mencion_poblacion_finita.group(0)
        else:
            # Por defecto, asumimos población infinita si no se menciona explícitamente
            return "infinita", "No se menciona población finita; se asume infinita."
    except Exception as e:
        print(f"Error al determinar el tamaño de la población: {e}")
        return "infinita", "Error al analizar el texto; se asume población infinita."

# Procesar texto y extraer parámetros del supuesto 1
def procesar_texto_supuesto1(texto):
    try:
        tasa_llegada = re.search(r'razón de\s*(\d+(\.\d+)?)\s*(pacientes por hora|clientes por minuto|clientes/hora|pacientes/minuto)', texto, re.IGNORECASE)
        tiempo_entre_llegadas = re.search(r'(tiempo esperado|media de|promedio de)\s*(\d+(\.\d+)?)\s*(minutos|horas)', texto, re.IGNORECASE)
        distribucion_tiempo = re.search(r'distribución de (Poisson|Exponencial)', texto, re.IGNORECASE)
        fragmento_tasa = tasa_llegada.group(0) if tasa_llegada else None
        fragmento_tiempo = tiempo_entre_llegadas.group(0) if tiempo_entre_llegadas else None

        if tasa_llegada:
            valor = tasa_llegada.group(1)
            unidades = tasa_llegada.group(3)
            fragmento = fragmento_tasa
            tipo = "tasa"
        elif tiempo_entre_llegadas:
            valor = tiempo_entre_llegadas.group(2)
            unidades = tiempo_entre_llegadas.group(4)
            fragmento = fragmento_tiempo
            tipo = "tiempo"
        else:
            valor, unidades, fragmento, tipo = None, None, "No encontrado", None

        return {
            "tipo": tipo,
            "valor": valor,
            "unidades": unidades,
            "distribucion": distribucion_tiempo.group(1) if distribucion_tiempo else "No especificada",
            "fragmento_enunciado": fragmento.strip()
        }
    except Exception as e:
        print(f"Error al procesar el texto para el supuesto 1: {e}")
        return None

import re

# Procesar texto y extraer parámetros del supuesto 2
def procesar_texto_supuesto2(texto):
    try:
        tasa_servicio = re.search(r'tasa de servicio de\s*(\d+(\.\d+)?)\s*(pacientes por hora|clientes por minuto|clientes/hora|pacientes/minuto)', texto, re.IGNORECASE)
        tiempo_servicio = re.search(r'(tiempo esperado|media de|promedio de)\s*(\d+(\.\d+)?)\s*(minutos|horas)', texto, re.IGNORECASE)

        # Buscar distribución exponencial
        distribucion_servicio = re.search(r'(distribución.*exponencial|tiempo.*es exponencial)', texto, re.IGNORECASE)

        # Ajustar texto de distribución
        distribucion_texto = "exponencial" if distribucion_servicio else "No especificada"

        fragmento_tasa = tasa_servicio.group(0) if tasa_servicio else None
        fragmento_tiempo = tiempo_servicio.group(0) if tiempo_servicio else None

        if tasa_servicio:
            valor = tasa_servicio.group(1)
            unidades = tasa_servicio.group(3)
            fragmento = fragmento_tasa
            tipo = "tasa"
        elif tiempo_servicio:
            valor = tiempo_servicio.group(2)
            unidades = tiempo_servicio.group(4)
            fragmento = fragmento_tiempo
            tipo = "tiempo"
        else:
            valor, unidades, fragmento, tipo = None, None, "No encontrado", None

        return {
            "tipo": tipo,
            "valor": valor,
            "unidades": unidades,
            "distribucion": distribucion_texto,
            "fragmento_enunciado": fragmento.strip() if fragmento else "No encontrado"
        }
    except Exception as e:
        print(f"Error al procesar el texto para el supuesto 2: {e}")
        return None


# Procesar texto y extraer parámetros del supuesto 5
def procesar_texto_supuesto5(texto):
    try:
        # Buscar menciones explícitas de servidores
        numero_servidores = re.search(r'(\d+)\s*servidores?\s*(en paralelo|simultáneos|disponibles)', texto, re.IGNORECASE)
        
        if numero_servidores:
            valor = numero_servidores.group(1)
            fragmento = numero_servidores.group(0)
        else:
            # Si no se menciona servidores, buscar contexto de un único médico o equivalente
            unico_servidor = re.search(r'clínica de un médico|consultorio de un médico|un médico', texto, re.IGNORECASE)
            if unico_servidor:
                valor = "1"  # Asumir un servidor por contexto
                fragmento = unico_servidor.group(0)
            else:
                valor = "1"  # Valor por defecto
                fragmento = "No se mencionaron servidores explícitos; se asume 1 servidor por defecto."
        
        return {
            "numero_servidores": int(valor),
            "fragmento_enunciado": fragmento.strip()
        }
    except Exception as e:
        print(f"Error al procesar el texto para el supuesto 5: {e}")
        return {"numero_servidores": 1, "fragmento_enunciado": "Error al determinar servidores, se asume 1 servidor por defecto."}

# Procesar texto y extraer parámetros del supuesto 6
def procesar_texto_supuesto6(texto, numero_servidores):
    try:
        # Buscar indicaciones de capacidad finita
        capacidad_sala = re.search(r'(sala de espera no puede acomodar más de)\s*(\d+)', texto, re.IGNORECASE)
        perdida_clientes = re.search(r'(pérdida de clientes|rechazo de clientes|no se puede acomodar más de)', texto, re.IGNORECASE)
        
        # Determinar capacidad total del sistema
        if capacidad_sala:
            capacidad_sala_num = int(capacidad_sala.group(2))  # Capacidad de la sala
            capacidad_total = capacidad_sala_num + numero_servidores  # Sumar servidores
            fragmento = capacidad_sala.group(0)
        elif perdida_clientes:
            # Inferir capacidad si se menciona pérdida de clientes y un número
            numero_pacientes = re.search(r'(\d+)\s*(pacientes|clientes)', texto, re.IGNORECASE)
            if numero_pacientes:
                capacidad_total = int(numero_pacientes.group(1)) + numero_servidores
                fragmento = perdida_clientes.group(0) + " " + numero_pacientes.group(0)
            else:
                capacidad_total = "No especificada"
                fragmento = perdida_clientes.group(0)
        else:
            capacidad_total = "Infinita"  # Por defecto, capacidad infinita
            fragmento = "No se mencionó ninguna limitación explícita de capacidad."

        return {
            "capacidad": capacidad_total,
            "fragmento_enunciado": fragmento.strip()
        }
    except Exception as e:
        print(f"Error al procesar el texto para el supuesto 6: {e}")
        return None

# Procesar texto y determinar si la población es finita o infinita
def procesar_texto_poblacion(texto):
    try:
        # Buscar indicaciones de población finita
        poblacion_fija = re.search(r'(población total|población finita|población limitada)\s*(de|es|para)\s*(\d+)', texto, re.IGNORECASE)
        limite_poblacion = re.search(r'(un máximo de|no más de|hasta)\s*(\d+)\s*(personas|clientes|pacientes)', texto, re.IGNORECASE)

        if poblacion_fija:
            valor = poblacion_fija.group(3)
            fragmento = poblacion_fija.group(0)
        elif limite_poblacion:
            valor = limite_poblacion.group(2)
            fragmento = limite_poblacion.group(0)
        else:
            valor = "Infinita"  # Por defecto, población infinita
            fragmento = "No se mencionó ninguna limitación explícita de población."

        return {
            "poblacion": valor,
            "fragmento_enunciado": fragmento.strip()
        }
    except Exception as e:
        print(f"Error al procesar el texto para la población: {e}")
        return None
    
# Crear Word para población finita
def manejar_poblacion_finita(doc, enunciado, fragmento_poblacion):
    doc.add_heading("Población Finita", level=1)
    doc.add_paragraph(f"✓ {enunciado}")
    doc.add_paragraph("✓ Es un tema más avanzado que no podrá solucionarse con este Programa.")
    doc.add_paragraph(f"✓ Texto del enunciado: \"{fragmento_poblacion}\".")

# Crear Word para población infinita
def manejar_poblacion_infinita(doc, parametros1, enunciado):
    doc.add_heading("Numeral 25: Continuación con Población Infinita", level=1)
    doc.add_paragraph(f"✓ {enunciado}")
    doc.add_paragraph(f"La distribución de probabilidad del tiempo es \"{parametros1['distribucion']}\".")
    if parametros1["tipo"] == "tasa":
        doc.add_paragraph(f"La tasa media de llegada es de {parametros1['valor']} {parametros1['unidades']}.")
    elif parametros1["tipo"] == "tiempo":
        doc.add_paragraph(f"El tiempo promedio entre llegadas es de {parametros1['valor']} {parametros1['unidades']}.")
    doc.add_paragraph(f"Texto del enunciado: \"{parametros1['fragmento_enunciado']}\".")

def identify_queue_discipline(text):
    discipline_keywords = {
        "FIFO": ["primero en llegar, primero en ser atendido", "first in, first out", "FIFO"],
        "LIFO": ["último en llegar, primero en ser atendido", "last in, first out", "LIFO"],
        "Priority": ["atención prioritaria", "priority queue", "prioridad", "prioritario"]
    }

    discipline = "No especificada"
    discipline_fragment = "No se especifica la disciplina de la cola."

    for disc, keywords in discipline_keywords.items():
        for keyword in keywords:
            if re.search(re.escape(keyword), text, re.IGNORECASE):
                discipline = disc
                discipline_fragment = keyword
                break
        if discipline != "No especificada":
            break

    return discipline, discipline_fragment


def crear_word_combinado_improved(
    parametros1, parametros2, parametros5, parametros6, parametros_poblacion, texto, ruta_destino
):
    doc = Document()
    
    # Supuesto 1
    doc.add_heading("Supuesto 1", level=1)
    doc.add_paragraph("Enunciado 1.")
    doc.add_paragraph(f"La distribución de probabilidad del tiempo es \"{parametros1['distribucion']}\".")
    if parametros1["tipo"] == "tasa":
        doc.add_paragraph(f"La tasa media de llegada es de {parametros1['valor']}.")
        doc.add_paragraph(f"Las unidades son {parametros1['unidades']}.")
    elif parametros1["tipo"] == "tiempo":
        doc.add_paragraph(f"El tiempo promedio entre llegadas es de {parametros1['valor']}.")
        doc.add_paragraph(f"Las unidades son {parametros1['unidades']}.")
    else:
        doc.add_paragraph("No se pudo determinar la tasa media de llegada o el tiempo promedio entre llegadas.")
    doc.add_paragraph(f"Texto del enunciado: \"{parametros1['fragmento_enunciado']}\".")
    
    # Supuesto 2
    doc.add_heading("Supuesto 2", level=1)
    doc.add_paragraph("Enunciado 1.")
    doc.add_paragraph(f"La distribución de probabilidad del tiempo es \"{parametros2['distribucion']}\".")
    if parametros2["tipo"] == "tasa":
        doc.add_paragraph(f"La tasa media de servicio es de {parametros2['valor']}.")
        doc.add_paragraph(f"Las unidades son {parametros2['unidades']}.")
    elif parametros2["tipo"] == "tiempo":
        doc.add_paragraph(f"El tiempo promedio de servicio es de {parametros2['valor']}.")
        doc.add_paragraph(f"Las unidades son {parametros2['unidades']}.")
    else:
        doc.add_paragraph("No se pudo determinar la tasa media de servicio o el tiempo promedio de servicio.")
    doc.add_paragraph(f"Texto del enunciado: \"{parametros2['fragmento_enunciado']}\".")
    
    # Supuesto 3 y 4
    doc.add_heading("Supuesto 3", level=1)
    doc.add_paragraph("Supuesto 3. La variable aleatoria del supuesto 1 (el tiempo que falta hasta el próximo nacimiento) y la variable aleatoria del supuesto 2 (el tiempo que falta hasta la siguiente muerte) son mutuamente independientes. La siguiente transición del estado del proceso es")
    doc.add_paragraph("n → n + 1 (un solo nacimiento)")
    doc.add_paragraph("o")
    doc.add_paragraph("n → n - 1 (una sola muerte),")
    doc.add_paragraph("lo que depende de cuál de las dos variables es más pequeña.")
    
    doc.add_heading("Supuesto 4", level=1)
    doc.add_paragraph("Supuesto 4. Se procederá cuando el sistema haya alcanzado la condición de estado estable (en caso de que pueda alcanzarla). Es decir, la tasa media a la que el proceso entra al estado n es igual a la tasa media a la que el proceso sale del estado n.")
    
    # Supuesto 5
    doc.add_heading("Supuesto 5", level=1)
    doc.add_paragraph("Enunciado 1.")
    doc.add_paragraph(f"El número de servidores en paralelo es \"{parametros5['numero_servidores']}\".")
    doc.add_paragraph(f"Texto del enunciado: \"{parametros5['fragmento_enunciado']}\".")
    
    # Supuesto 6
    if parametros6["capacidad"] != "Infinita":
        doc.add_heading("Supuesto 6 capacidad finita", level=1)
        doc.add_paragraph("Enunciado 1.")
        doc.add_paragraph(f"La capacidad del sistema es de \"{parametros6['capacidad']}\".")
        doc.add_paragraph(f"Texto del enunciado: \"{parametros6['fragmento_enunciado']}\".")
    
    # Población
    if parametros_poblacion["poblacion"] != "Infinita":
        manejar_poblacion_finita(doc, "Enunciado 1", parametros_poblacion["fragmento_enunciado"])
    else:
        # Numeral 25 & 26: Queue Discipline
        discipline, discipline_fragment = identify_queue_discipline(texto)
        doc.add_heading("Disciplina de atención", level=1)
        doc.add_paragraph(f"✓ Enunciado 1.")
        doc.add_paragraph(f"✓ La disciplina del sistema es de “{discipline}”.")
        doc.add_paragraph(f"✓ Texto del enunciado: “{discipline_fragment}”.")


    # Numeral 27: Check for "Enunciado 1" and add to "Supuestos de Enunciados Distintos"
    enunciado1_found = False
    for paragraph in doc.paragraphs:
        if "Enunciado 1" in paragraph.text and ("NO proceso nacimiento muerte" in paragraph.text or "Población Finita" in paragraph.text):
            enunciado1_found = True
            break

    if enunciado1_found:
        doc.add_heading("Supuestos de Enunciados Distintos", level=1)
        doc.add_paragraph(f"✓ Enunciado 1")
        doc.add_paragraph("✓ Es un tema más avanzado que no podrá solucionarse con este Programa.")

    # Guardar archivo
    doc.save(ruta_destino)
    print(f"Documento generado: {ruta_destino}")

# Ruta de los archivos Word
ruta_word = input("\nPor favor para comenzar, ingresa la ruta del archivo Word: ")
ruta_destino = r"WORD #2.docx"  # Reemplaza con la ruta del archivo Word destino

# Proceso completo
texto_transcrito = leer_texto_desde_word(ruta_word)
parametros_supuesto1 = procesar_texto_supuesto1(texto_transcrito)
parametros_supuesto2 = procesar_texto_supuesto2(texto_transcrito)
# Extraer parámetros de los textos para los supuestos 5 y 6
parametros_supuesto5 = procesar_texto_supuesto5(texto_transcrito)

if parametros_supuesto5:  # Validar que se procesó correctamente
    numero_servidores = parametros_supuesto5["numero_servidores"]
    parametros_supuesto6 = procesar_texto_supuesto6(texto_transcrito, numero_servidores)
else:
    print("Error al procesar el texto para el supuesto 5.")
    numero_servidores = 1  # Valor por defecto
    parametros_supuesto6 = procesar_texto_supuesto6(texto_transcrito, numero_servidores)

# Procesar población
parametros_poblacion = procesar_texto_poblacion(texto_transcrito)

# Crear el archivo Word con la información combinada
crear_word_combinado_improved(
    parametros1=parametros_supuesto1,
    parametros2=parametros_supuesto2,
    parametros5=parametros_supuesto5,
    parametros6=parametros_supuesto6,
    parametros_poblacion=parametros_poblacion,
    texto=texto_transcrito,
    ruta_destino=ruta_destino
)