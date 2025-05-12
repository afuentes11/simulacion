import re
from PIL import Image, ImageEnhance
import pytesseract
from docx import Document

# Configura la ruta a Tesseract si estás en Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Preprocesamiento de imagen para mejorar la transcripción
def preprocesar_imagen(ruta_imagen):
    try:
        imagen = Image.open(ruta_imagen)
        imagen = imagen.convert('L')  # Escala de grises
        enhancer = ImageEnhance.Contrast(imagen)
        imagen = enhancer.enhance(2)  # Mejorar contraste
        imagen = imagen.point(lambda x: 255 if x > 138 else 0, '1')  # Binarización
        return imagen
    except Exception as e:
        raise ValueError(f"Error al preprocesar la imagen: {e}")

# Transcribir texto desde imagen
def transcribir_imagen(ruta_imagen):
    try:
        imagen = preprocesar_imagen(ruta_imagen)
        texto = pytesseract.image_to_string(imagen, lang='spa')
        return texto
    except Exception as e:
        return f"Error al procesar la imagen: {e}"

# Guardar texto en un archivo de Word
def guardar_en_word(texto, ruta_word):
    try:
        documento = Document()
        documento.add_heading("Enunciado 1", level=1)
        documento.add_paragraph(texto)
        documento.save(ruta_word)
        print(f"El texto ha sido guardado en el archivo: {ruta_word}")
    except Exception as e:
        print(f"Error al guardar el archivo Word: {e}")

# Ruta de la imagen
ruta_imagen = input("Por favor para comenzar, ingresa la ruta de la imagen del enunciado: ")
# Ruta para guardar el archivo Word
ruta_word = r"WORD #1.docx"

# Paso 1: Transcribir el texto desde la imagen
texto_transcrito = transcribir_imagen(ruta_imagen)

# Paso 2: Guardar en un archivo Word
guardar_en_word(texto_transcrito, ruta_word)
