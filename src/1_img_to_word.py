import re
from PIL import Image, ImageEnhance
import easyocr
from docx import Document

# Inicializar el lector de EasyOCR (idiomas: español e inglés; gpu=False)
reader = easyocr.Reader(['es', 'en'], gpu=False)

# Preprocesamiento de imagen para mejorar la transcripción
def preprocesar_imagen(ruta_imagen):
    try:
        imagen = Image.open(ruta_imagen)
        imagen = imagen.convert('L')  # Escala de grises
        enhancer = ImageEnhance.Contrast(imagen)
        imagen = enhancer.enhance(2)  # Mejorar contraste
        imagen = imagen.point(lambda x: 255 if x > 138 else 0, '1')  # Binarización
        return imagen
    except Exception as e:
        raise ValueError(f"Error al preprocesar la imagen: {e}")

# Transcribir texto desde imagen usando EasyOCR
def transcribir_imagen(ruta_imagen):
    try:
        # EasyOCR puede aceptar ruta de archivo directamente, 
        # pero aplicamos preprocesamiento para mejorar precisión
        imagen_pre = preprocesar_imagen(ruta_imagen)
        # Convertir la imagen preprocesada a array para EasyOCR
        resultado = reader.readtext(
            # EasyOCR admite PIL.Image directamente en versiones recientes:
            imagen_pre, 
            detail=0,       # solo el texto (sin bounding boxes)
            paragraph=True  # agrupar en párrafos si es posible
        )
        # Unir todos los fragmentos de texto en un único bloque
        texto = "\n".join(resultado)
        return texto
    except Exception as e:
        return f"Error al procesar la imagen: {e}"

# Guardar texto en un archivo de Word
def guardar_en_word(texto, ruta_word):
    try:
        documento = Document()
        documento.add_heading("Enunciado 1", level=1)
        documento.add_paragraph(texto)
        documento.save(ruta_word)
        print(f"El texto ha sido guardado en el archivo: {ruta_word}")
    except Exception as e:
        print(f"Error al guardar el archivo Word: {e}")

# Rutas de trabajo
ruta_imagen = "EJERCICIO 14-MM1K.PNG"
ruta_word = "WORD #1.docx"

# Paso 1: Transcribir el texto desde la imagen
texto_transcrito = transcribir_imagen(ruta_imagen)

# Paso 2: Guardar en un archivo Word
guardar_en_word(texto_transcrito, ruta_word)

